"""
utils/resume_parser.py
──────────────────────
Extracts plain text from uploaded resume files (PDF or DOCX).
The extracted text is stored in the session and injected into the LLM prompt.

Improvements over v1:
  • PDF: uses get_text("blocks") to preserve layout across multi-column resumes,
    then sorts blocks by (Y, X) so reading order is correct.
  • DOCX: also reads section headers, footers, and text-box shapes (where
    name/contact info often lives in modern resume templates).
  • Cleaning: strips junk unicode, collapses excessive whitespace, deduplicates
    blank lines, and logs a preview so you can confirm what the LLM receives.
"""

import io
import re
import unicodedata
from pathlib import Path
from utils.logger import get_logger

logger = get_logger(__name__)


# ── PDF ───────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz).

    Uses the 'blocks' extraction mode which keeps each text block as a unit,
    then sorts by vertical position so multi-column PDFs are read top-to-bottom
    rather than jumbling the columns together.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.error("PyMuPDF not installed. Run: pip install pymupdf")
        return ""

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text: list[str] = []

        for page_num, page in enumerate(doc):
            # get_text("blocks") returns:
            # (x0, y0, x1, y1, text, block_no, block_type)
            # block_type == 0 → text block; 1 → image
            blocks = page.get_text("blocks")

            # Sort by top-left corner: primarily Y (row), secondarily X (column)
            text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
            text_blocks.sort(key=lambda b: (round(b[1] / 10), b[0]))  # bucket Y into ~10pt rows

            page_lines = []
            for b in text_blocks:
                block_text = b[4].strip()
                if block_text:
                    page_lines.append(block_text)

            if page_lines:
                pages_text.append("\n".join(page_lines))

        doc.close()
        full_text = "\n\n".join(pages_text)
        logger.info(
            f"PDF parsed: {len(full_text)} chars across {len(pages_text)} pages | "
            f"preview: {full_text[:200]!r}"
        )
        return full_text

    except Exception as e:
        logger.error(f"PDF parse error: {e}")
        return ""


# ── DOCX ──────────────────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Reads:
      • Body paragraphs (styled and unstyled)
      • Table cells (skills grids, education tables etc.)
      • Section headers and footers (where contact info often lives)
      • Text boxes / shapes in the document XML
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""

    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        # 1. Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 2. Tables (skills, education grids, etc.)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        # 3. Headers and footers (contact info, name, phone often here)
        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf is not None:
                    for para in hf.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)

        # 4. Text boxes / drawing shapes in body XML
        # These appear in modern resume templates (floating name blocks etc.)
        body = doc.element.body
        for txbx in body.iter(qn("w:txbxContent")):
            for p in txbx.iter(qn("w:p")):
                text = "".join(
                    r.text for r in p.iter(qn("w:t")) if r.text
                ).strip()
                if text:
                    parts.append(text)

        full_text = "\n".join(parts)
        logger.info(
            f"DOCX parsed: {len(full_text)} chars | "
            f"preview: {full_text[:200]!r}"
        )
        return full_text

    except Exception as e:
        logger.error(f"DOCX parse error: {e}")
        return ""


# ── Entry point ───────────────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Auto-detect file type and extract resume text.

    Args:
        file_bytes: Raw file content
        filename:   Original filename (used to detect extension)

    Returns:
        Cleaned plain text, or empty string on failure
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        text = extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
    else:
        logger.warning(f"Unsupported resume format: {ext}")
        return ""

    return _clean_text(text)


# ── Text cleaner ──────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Normalise extracted resume text:
      • Normalise unicode (NFC) and remove non-printable control characters
      • Collapse ligatures and unusual whitespace
      • Strip leading/trailing whitespace per line
      • Allow at most one consecutive blank line between sections
      • Remove lines that are just punctuation noise (e.g. "---", "•••")
    """
    if not text:
        return ""

    # 1. Unicode normalisation — fix ligatures (ﬁ → fi) and composed chars
    text = unicodedata.normalize("NFKC", text)

    # 2. Remove non-printable / control characters (keep \n and \t)
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)

    # 3. Collapse inline runs of whitespace (not newlines) to single space
    text = re.sub(r"[ \t\u00A0]+", " ", text)

    # 4. Process line by line
    lines = text.splitlines()
    cleaned: list[str] = []
    blank_count = 0

    for line in lines:
        stripped = line.strip()

        # Skip lines that are purely decorative (e.g. "────", "====", "• • •")
        if stripped and re.fullmatch(r"[-=_*•·|/\\]{3,}", stripped):
            continue

        if stripped:
            cleaned.append(stripped)
            blank_count = 0
        else:
            blank_count += 1
            if blank_count <= 1:   # allow max 1 blank line between sections
                cleaned.append("")

    return "\n".join(cleaned).strip()


def extract_candidate_name(resume_text: str) -> str:
    """
    Heuristic: the candidate's name is almost always the first non-empty,
    short line of the resume (< 50 chars, no '@', no digits).

    Returns an empty string if no suitable name line is found.
    """
    if not resume_text:
        return ""

    for line in resume_text.splitlines():
        line = line.strip()
        if not line:
            continue
        # Skip lines that look like contact info, URLs, or headers
        if len(line) > 50:
            continue
        if any(c in line for c in ("@", "http", "www", "|", "linkedin")):
            continue
        # Must contain at least two words (first + last name)
        words = line.split()
        if len(words) < 2:
            continue
        # Skip lines that are all-caps section headers like "WORK EXPERIENCE"
        if line.isupper() and len(words) <= 4:
            continue
        logger.info(f"Extracted candidate name: '{line}'")
        return line

    return ""
